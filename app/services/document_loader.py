from __future__ import annotations

import csv
from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

from app.core.manual_routing import get_manual_family


SUPPORTED_EXTENSIONS = {".pdf", ".csv", ".docx"}


def _manual_metadata(file_path: Path) -> dict[str, str]:
    stem = file_path.stem
    code = stem.split("_", 1)[0].upper()
    return {
        "source": file_path.name,
        "document_title": stem.replace("_", " "),
        "manual_code": code,
        "manual_family": get_manual_family(code),
    }


def _load_pdf_document(file_path: Path) -> list[Document]:
    loader = PyPDFLoader(str(file_path))
    pages = loader.load()
    for page in pages:
        page.metadata.update(_manual_metadata(file_path))
        page.metadata["file_type"] = "pdf"
    return pages


def _load_csv_document(file_path: Path) -> list[Document]:
    documents: list[Document] = []

    with file_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        if not reader.fieldnames:
            raise ValueError(f"El archivo CSV no tiene encabezados: {file_path.name}")

        for row_number, row in enumerate(reader, start=1):
            row_content = "\n".join(
                f"{field}: {value}"
                for field, value in row.items()
                if value not in (None, "")
            ).strip()

            if not row_content:
                continue

            documents.append(
                Document(
                    page_content=row_content,
                    metadata={
                        **_manual_metadata(file_path),
                        "page": row_number,
                        "row": row_number,
                        "file_type": "csv",
                    },
                )
            )

    if not documents:
        raise ValueError(f"El archivo CSV no contiene filas útiles: {file_path.name}")

    return documents


def _load_docx_document(file_path: Path) -> list[Document]:
    docx_file = DocxDocument(str(file_path))
    paragraph_texts = [paragraph.text.strip() for paragraph in docx_file.paragraphs if paragraph.text.strip()]

    table_texts: list[str] = []
    for table in docx_file.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if table_rows:
            table_texts.append("\n".join(table_rows))

    content_parts = paragraph_texts + table_texts
    content = "\n\n".join(part for part in content_parts if part).strip()
    if not content:
        raise ValueError(f"El archivo DOCX no contiene texto útil: {file_path.name}")

    return [
        Document(
            page_content=content,
            metadata={
                **_manual_metadata(file_path),
                "page": 1,
                "file_type": "docx",
            },
        )
    ]


def _load_supported_document(file_path: Path) -> list[Document]:
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return _load_pdf_document(file_path)

    if extension == ".csv":
        return _load_csv_document(file_path)

    if extension == ".docx":
        return _load_docx_document(file_path)

    raise ValueError(f"Formato no soportado: {file_path.name}")


def load_documents_from_path(docs_path: str) -> list[Document]:
    base_path = Path(docs_path)
    if not base_path.exists():
        raise FileNotFoundError(f"La ruta de documentos no existe: {docs_path}")

    supported_files = [
        file_path
        for file_path in sorted(base_path.iterdir())
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    if not supported_files:
        raise FileNotFoundError(
            f"No se encontraron archivos PDF, CSV o DOCX en: {docs_path}"
        )

    documents: list[Document] = []
    for file_path in supported_files:
        documents.extend(_load_supported_document(file_path))

    return documents


def load_documents_from_files(file_paths: list[Path]) -> list[Document]:
    documents: list[Document] = []

    for file_path in sorted(file_paths):
        if not file_path.exists() or not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        documents.extend(_load_supported_document(file_path))

    if not documents:
        raise FileNotFoundError("No se encontraron documentos válidos para indexar")

    return documents


def load_pdf_documents(docs_path: str) -> list[Document]:
    return load_documents_from_path(docs_path)
